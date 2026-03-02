import asyncio
from pathlib import Path
from rag_forge.document.models import Document
from rag_forge.core.interfaces import DocumentLoader

class PDFLoader:
    async def load(self, path: str) -> list[Document]:
        import pymupdf4llm, pymupdf

        file_path = Path(path)
        doc = pymupdf.open(file_path)
        total_pages = doc.page_count
        content = pymupdf4llm.to_markdown(doc)
        source = file_path.name
        file_type = "pdf"
        metadata = {"total_pages": total_pages}
        return [Document(content=content, source=source, file_type=file_type, metadata=metadata)]

class DOCXLoader:
    async def load(self, path: str) -> list[Document]:
        import docx as python_docx
        file_path = Path(path)
        doc = python_docx.Document(file_path)
        content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        source = file_path.name
        file_type = "docx"
        return [Document(content=content, source=source, file_type=file_type)]

class TXTLoader:
    async def load(self, path: str) -> list[Document]:
        file_path = Path(path)
        loop = asyncio.get_running_loop()
        content = await loop.run_in_executor(None, file_path.read_text)
        source = file_path.name
        file_type = "txt"
        return [Document(content=content, source=source, file_type=file_type)]

class DocumentLoaderFactory:
    _LOADERS: dict[str, type[DocumentLoader]] = {
        ".pdf": PDFLoader,
        ".docx": DOCXLoader,
        ".txt": TXTLoader
    }

    @classmethod
    def get_loader(cls, path: str) -> DocumentLoader:
        file_path = Path(path)
        file_ext = file_path.suffix.lower()
        if file_ext in cls._LOADERS:
            return cls._LOADERS[file_ext]()
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")




